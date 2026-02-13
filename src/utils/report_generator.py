from docx import Document
from docx.shared import Inches
import os


def generate_report(analysis: dict, resume_text: str, jd_text: str) -> str:
    """Generate a DOCX report from the analysis."""
    document = Document()
    document.add_heading("CV Embed Analysis Report", 0)

    document.add_heading("Overall Summary", level=1)
    document.add_paragraph(analysis.get("overall_summary", "No summary available."))

    document.add_heading("Experience Analysis", level=1)
    document.add_paragraph(
        analysis.get("experience_analysis", "No experience analysis available.")
    )

    document.add_heading("Skill Analysis", level=1)
    document.add_paragraph(
        analysis.get("skill_analysis", "No skill analysis available.")
    )

    document.add_heading("Education Analysis", level=1)
    document.add_paragraph(
        analysis.get("education_analysis", "No education analysis available.")
    )

    document.add_heading("Suitability Score", level=1)
    document.add_paragraph(f"{analysis.get('suitability_score', 'N/A')}/100")

    document.add_heading("Report Summary", level=1)
    report_summary = analysis.get("report_summary", [])
    if isinstance(report_summary, list):
        for item in report_summary:
            document.add_paragraph(item, style="List Bullet")
    else:
        document.add_paragraph(str(report_summary))

    document.add_heading("Resume", level=1)
    document.add_paragraph(resume_text)

    document.add_heading("Job Description", level=1)
    document.add_paragraph(jd_text)

    # Ensure the uploads directory exists
    if not os.path.exists("uploads"):
        os.makedirs("uploads")

    report_path = os.path.join("uploads", "report.docx")
    document.save(report_path)
    return report_path
